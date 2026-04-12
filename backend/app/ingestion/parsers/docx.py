from pathlib import Path
from docx import Document as DocxDocument
from app.ingestion.parsers.base import BaseParser, ParsedPage, ParseResult

class DocxParser(BaseParser):
    supported_extensions = [".docx"]

    def parse(self, file_path: Path) -> ParseResult:
        doc = DocxDocument(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))
        full_text = "\n\n".join(paragraphs)
        metadata = {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        if doc.core_properties:
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
        return ParseResult(pages=[ParsedPage(text=full_text, page_number=1)], metadata=metadata, file_type="docx")
